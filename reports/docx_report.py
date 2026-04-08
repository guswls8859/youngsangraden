"""
용산어린이정원 일일업무보고 - python-docx (.docx) 생성 모듈
"""
import io
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── 헬퍼 ────────────────────────────────────────────────────────────────────

def _set_bg(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _ct(cell, text, bold=False, align='left', bg=None, font_size=9):
    """셀에 텍스트를 설정한다."""
    FN = '맑은 고딕'
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.alignment = {
        'left':   WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right':  WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(str(text) if text is not None else '')
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.name = FN
    run._r.rPr.rFonts.set(qn('w:eastAsia'), FN)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if bg:
        _set_bg(cell, bg)


def _fmt(n, suffix=''):
    try:
        return f'{int(n):,}{suffix}'
    except (TypeError, ValueError):
        return f'0{suffix}'


def _disp(v, suffix=''):
    """None → '-', 숫자 → 포맷된 문자열"""
    if v is None:
        return '-'
    try:
        return f'{int(v):,}{suffix}'
    except (TypeError, ValueError):
        return str(v)


# ── 메인 함수 ────────────────────────────────────────────────────────────────

def build_integrated_daily_docx(target_date, ops, sf_slots,
                                 eoulrim, jamjam, kumnare, info_report,
                                 info_shuttle_items=None, info_patrol_items=None,
                                 total_sales=0):
    doc = Document()

    # 페이지 설정 (A4)
    sec = doc.sections[0]
    sec.page_width    = Mm(210)
    sec.page_height   = Mm(297)
    sec.left_margin   = Mm(18)
    sec.right_margin  = Mm(18)
    sec.top_margin    = Mm(18)
    sec.bottom_margin = Mm(18)

    # ── 제목 ──────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after  = Pt(2)
    run = title_p.add_run(f'{target_date.strftime("%Y.%m.%d.")}  (운영관리) 용산어린이정원 일일업무보고')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = '맑은 고딕'
    run._r.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    # ── 결재란 ───────────────────────────────────────────
    ct = doc.add_table(rows=2, cols=4)
    ct.style = 'Table Grid'
    for i, lbl in enumerate(['담당', '운영1팀장', '운영2팀장', '시설관리팀장']):
        _ct(ct.rows[0].cells[i], lbl, bold=True, align='center', bg='F0F0F0', font_size=8)
        _ct(ct.rows[1].cells[i], '', font_size=20)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── 메인 보고 테이블 ─────────────────────────────────
    tbl = doc.add_table(rows=0, cols=5)
    tbl.style = 'Table Grid'
    # 열 너비: 섹션번호 | 섹션명 | 서브레이블 | 내용A | 내용B
    # 총 170mm
    for i, w in enumerate([Mm(14), Mm(20), Mm(20), Mm(58), Mm(58)]):
        tbl.columns[i].width = w

    def _row_5(vals, bgs=None, bolds=None, aligns=None, fsizes=None):
        """5열 행 추가"""
        row = tbl.add_row()
        for i, v in enumerate(vals):
            bg    = bgs[i]    if bgs    else None
            bold  = bolds[i]  if bolds  else False
            align = aligns[i] if aligns else 'left'
            fs    = fsizes[i] if fsizes  else 9
            _ct(row.cells[i], v, bold=bold, align=align, bg=bg, font_size=fs)
        return row

    # ① 금일 방문현황
    row = tbl.add_row()
    row.cells[0].merge(row.cells[1])
    _ct(row.cells[0], '■\n금일\n방문\n현황', bold=True, align='center', bg='1F3864', font_size=8)
    _set_bg(row.cells[0], '1F3864')
    # 폰트 색 흰색
    for p in row.cells[0].paragraphs:
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    row.cells[2].merge(row.cells[4])
    _ct(row.cells[2],
        f'입장  {_fmt(ops.today_total if ops else 0)}명\n'
        f'   주출입구 도보방문(명): {_fmt(ops.main_gate_walk if ops else 0)}   '
        f'부출입구 도보방문(명): {_fmt(ops.sub_gate_walk if ops else 0)}   '
        f'차량방문(명): {_fmt(ops.car_visit if ops else 0)}',
        font_size=9)

    # ② 전일 방문현황 / 명일 기상
    row = tbl.add_row()
    row.cells[0].merge(row.cells[1])
    _ct(row.cells[0], '■\n전일\n방문\n현황', bold=True, align='center', bg='1F3864', font_size=8)
    _set_bg(row.cells[0], '1F3864')
    for p in row.cells[0].paragraphs:
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    row.cells[2].merge(row.cells[3])
    _ct(row.cells[2],
        f'입장  {_fmt(ops.yesterday_total if ops else 0)}명',
        font_size=10, bold=True)
    _ct(row.cells[4],
        f'명일 기상상황\n기온 {ops.tomorrow_temp_min if ops else 0}°~ {ops.tomorrow_temp_max if ops else 0}°\n강수확률 {ops.tomorrow_rain_pct if ops else 0}%',
        font_size=9)

    # ③ 운영관리 - 내부시설
    row = tbl.add_row()
    row.cells[0].merge(row.cells[0])  # no merge yet
    _ct(row.cells[0], '■\n운영\n관리', bold=True, align='center', bg='1F3864', font_size=8)
    _set_bg(row.cells[0], '1F3864')
    for p in row.cells[0].paragraphs:
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    _ct(row.cells[1], '시설', bold=True, align='center', bg='DCE6F1', font_size=8)
    _ct(row.cells[2], '내부시설\n(용산서가, 전시관 등)', bold=False, align='center', bg='F2F2F2', font_size=7.5)
    row.cells[3].merge(row.cells[4])
    _ct(row.cells[3], getattr(ops, 'facility_interior', '') or '' if ops else '', font_size=8.5)

    # 잔디마당 - 운영관리 rowspan 처리: 같은 첫 번째 셀 재사용 불가 (docx 방식)
    for lbl, attr in [
        ('잔디마당\n가로수길\n전망언덕', 'facility_outdoor'),
        ('분수정원\n잼잼카페',           'facility_fountain'),
    ]:
        row = tbl.add_row()
        _ct(row.cells[0], '', bg='1F3864')
        _set_bg(row.cells[0], '1F3864')
        _ct(row.cells[1], '', bg='DCE6F1')
        _ct(row.cells[2], lbl, align='center', bg='F2F2F2', font_size=7.5)
        row.cells[3].merge(row.cells[4])
        _ct(row.cells[3], getattr(ops, attr, '') or '' if ops else '', font_size=8.5)

    # 스포츠필드 (내부 테이블은 텍스트로 단순 표현)
    row = tbl.add_row()
    _ct(row.cells[0], '', bg='1F3864')
    _set_bg(row.cells[0], '1F3864')
    _ct(row.cells[1], '', bg='DCE6F1')
    _ct(row.cells[2], '스포츠\n필드', align='center', bg='F2F2F2', font_size=7.5)
    row.cells[3].merge(row.cells[4])
    sf_cell = row.cells[3]
    sf_cell.text = ''

    # 스포츠필드 내부 소형 테이블
    sf_inner = sf_cell.add_table(rows=1, cols=7)
    sf_inner.style = 'Table Grid'
    for i, w in enumerate([Mm(22), Mm(10), Mm(10), Mm(10), Mm(10), Mm(10), Mm(10)]):
        sf_inner.columns[i].width = w

    def _sf_hdr(cells, labels, bg='D9E1F2'):
        for i, lbl in enumerate(labels):
            _ct(cells[i], lbl, bold=True, align='center', bg=bg, font_size=7)

    # 축구장/테니스장 헤더
    h1 = sf_inner.rows[0]
    h1.cells[0].merge(h1.cells[0])
    _ct(h1.cells[0], '구분', bold=True, align='center', bg='D9E1F2', font_size=7)
    h1.cells[1].merge(h1.cells[3])
    _ct(h1.cells[1], '축구장', bold=True, align='center', bg='D9E1F2', font_size=7)
    h1.cells[4].merge(h1.cells[6])
    _ct(h1.cells[4], '테니스장', bold=True, align='center', bg='D9E1F2', font_size=7)

    sub1 = sf_inner.add_row()
    _ct(sub1.cells[0], '', bg='D9E1F2')
    for i, l in enumerate(['분류', '예약', '입장', '분류', '예약', '입장'], 1):
        _ct(sub1.cells[i], l, bold=True, align='center', bg='EBEBEB', font_size=7)

    st_rows = sf_slots.get('st_rows', []) if sf_slots else []
    for sr in st_rows:
        r = sf_inner.add_row()
        _ct(r.cells[0], sr['label'], font_size=7)
        _ct(r.cells[1], sr['soccer'].get('cat') or '-', align='center', font_size=7)
        _ct(r.cells[2], _disp(sr['soccer'].get('reserved')), align='center', font_size=7)
        _ct(r.cells[3], _disp(sr['soccer'].get('actual')),   align='center', font_size=7)
        _ct(r.cells[4], sr['tennis'].get('cat') or '-', align='center', font_size=7)
        _ct(r.cells[5], _disp(sr['tennis'].get('reserved')), align='center', font_size=7)
        _ct(r.cells[6], _disp(sr['tennis'].get('actual')),   align='center', font_size=7)

    # 야구장/합계 헤더
    h2 = sf_inner.add_row()
    _ct(h2.cells[0], '구분', bold=True, align='center', bg='D9E1F2', font_size=7)
    h2.cells[1].merge(h2.cells[3])
    _ct(h2.cells[1], '야구장', bold=True, align='center', bg='D9E1F2', font_size=7)
    h2.cells[4].merge(h2.cells[6])
    _ct(h2.cells[4], '합계', bold=True, align='center', bg='D9E1F2', font_size=7)

    sub2 = sf_inner.add_row()
    _ct(sub2.cells[0], '', bg='D9E1F2')
    for i, l in enumerate(['분류', '예약', '입장', '분류', '예약', '입장'], 1):
        _ct(sub2.cells[i], l, bold=True, align='center', bg='EBEBEB', font_size=7)

    bb_rows = sf_slots.get('bb_rows', []) if sf_slots else []
    for br in bb_rows:
        r = sf_inner.add_row()
        _ct(r.cells[0], br['label'], font_size=7)
        _ct(r.cells[1], br['baseball'].get('cat') or '-', align='center', font_size=7)
        _ct(r.cells[2], _disp(br['baseball'].get('reserved')), align='center', font_size=7)
        _ct(r.cells[3], _disp(br['baseball'].get('actual')),   align='center', font_size=7)
        _ct(r.cells[4], br['total'].get('cat') or '-', align='center', font_size=7)
        _ct(r.cells[5], _disp(br['total'].get('reserved')), align='center', font_size=7)
        _ct(r.cells[6], _disp(br['total'].get('actual')),   align='center', font_size=7)

    # 주차장
    row = tbl.add_row()
    _ct(row.cells[0], '', bg='1F3864')
    _set_bg(row.cells[0], '1F3864')
    _ct(row.cells[1], '', bg='DCE6F1')
    _ct(row.cells[2], '주차장', align='center', bg='F2F2F2', font_size=7.5)
    row.cells[3].merge(row.cells[4])
    _ct(row.cells[3],
        f'다둥이 {_fmt(ops.parking_family if ops else 0)}대   '
        f'장애인 {_fmt(ops.parking_disabled if ops else 0)}대   '
        f'임산부 {_fmt(ops.parking_pregnant if ops else 0)}대   '
        f'어린이단체 {_fmt(ops.parking_children if ops else 0)}대',
        font_size=9)

    # 편익시설 매출
    e_s = eoulrim.daily_net_sales if eoulrim else 0
    j_s = jamjam.daily_net_sales  if jamjam  else 0
    k_s = kumnare.sales_amount    if kumnare  else 0
    row = tbl.add_row()
    row.cells[0].merge(row.cells[1])
    _ct(row.cells[0], '편익시설\n매출', bold=True, align='center', bg='E2EFDA', font_size=8)
    _ct(row.cells[2], '카페어울림', align='center', bg='E2EFDA', font_size=8)
    _ct(row.cells[3], _fmt(e_s, ' 원') + '  /  잼잼카페: ' + _fmt(j_s, ' 원') + '  /  꿈나래마켓: ' + _fmt(k_s, ' 원'),
        font_size=9)
    row.cells[4].merge(row.cells[4])
    _ct(row.cells[4], f'합계: {_fmt(total_sales, " 원")}', font_size=9)

    # 내부행사/외부행사
    for lbl, attr in [('내부행사\n/프로그램', 'internal_event'), ('외부행사', 'external_event')]:
        row = tbl.add_row()
        row.cells[0].merge(row.cells[1])
        _ct(row.cells[0], lbl, bold=True, align='center', font_size=8)
        row.cells[2].merge(row.cells[4])
        _ct(row.cells[2], getattr(ops, attr, '') or '' if ops else '', font_size=8.5)

    # 특이사항
    row = tbl.add_row()
    row.cells[0].merge(row.cells[1])
    _ct(row.cells[0], '특이사항', bold=True, align='center', font_size=8)
    row.cells[2].merge(row.cells[4])
    _ct(row.cells[2], getattr(ops, 'special_notes', '') or '' if ops else '', font_size=8.5)

    # 세부 이용현황
    row = tbl.add_row()
    row.cells[0].merge(row.cells[1])
    _ct(row.cells[0], '', font_size=8)
    _ct(row.cells[2], '○ 세부 이용현황\n셔틀버스: ' + _fmt(info_report.shuttle_total if info_report else 0, '명') +
        '  /  꿈나래마켓 대여물품: ' + _fmt(kumnare.rental_total_users if kumnare else 0, '명') +
        '  /  스탬프투어: ' + _fmt(kumnare.stamp_issued if kumnare else 0, '명'),
        font_size=9)
    row.cells[3].merge(row.cells[4])

    # 비상연락
    row = tbl.add_row()
    row.cells[0].merge(row.cells[1])
    _ct(row.cells[0], '비상\n연락', bold=True, align='center', font_size=8)
    _ct(row.cells[2], '운영대행', bold=True, align='center', bg='F2F2F2', font_size=7.5)
    _ct(row.cells[3],
        '(총괄) 이정환 팀장 010-6395-8035\n(운영) 김인경 과장 010-2039-0827',
        font_size=7.5)
    _ct(row.cells[4],
        'LH\n(운영총괄) 이원영 차장 010-9047-0752\n(대외협의) 김수민 과장 010-3928-8627',
        font_size=7.5)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
